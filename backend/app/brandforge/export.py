"""Export the generated Brand Document to markdown, docx, and a condensed
voice-system-prompt card. Ported from the BrandForge Space (src/export.py).
python-docx only.
"""
from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..services.markdown_docx import HEADING_RE, render_markdown_body
from .archetypes import ARCHETYPES_BY_ID
from .imaging import extract_palette, palette_markdown, strip_palette_block
from .intake import BrandIntake
from .sections import PHASES

VERSION = "1.0"


def _display_content(content: str) -> str:
    """Client-facing rendering of a section: the machine-readable ```palette
    block (Visual Direction Brief) becomes a readable list."""
    swatches = extract_palette(content)
    if not swatches:
        return content
    return strip_palette_block(content) + "\n\n" + palette_markdown(swatches)


def _front_matter(intake: BrandIntake) -> str:
    return (
        "---\n"
        f"brand: {intake.brand_name}\n"
        f"category: {intake.brand_category}\n"
        f"date: {date.today().isoformat()}\n"
        f"version: {VERSION}\n"
        "---\n\n"
    )


def to_markdown(intake: BrandIntake, sections: dict[str, str]) -> str:
    parts = [_front_matter(intake), f"# {intake.brand_name} — Brand Document\n"]
    for phase, names in PHASES.items():
        phase_parts = []
        for name in names:
            content = _display_content(sections.get(name, "").strip())
            if not content:
                continue
            if not content.lstrip().startswith("#"):
                phase_parts.append(f"## {name}\n\n{content}\n")
            else:
                phase_parts.append(f"{content}\n")
        if phase_parts:
            parts.append(f"# {phase}\n")
            parts.extend(phase_parts)
    return "\n".join(parts)


# --- docx ----------------------------------------------------------------


def to_docx_bytes(intake: BrandIntake, sections: dict[str, str]) -> bytes:
    doc = Document()

    title = doc.add_heading(intake.brand_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Brand Document")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    meta = doc.add_paragraph(f"{intake.brand_category}  |  {date.today().isoformat()}  |  v{VERSION}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    for phase, names in PHASES.items():
        written = [n for n in names if sections.get(n, "").strip()]
        if not written:
            continue
        doc.add_paragraph(phase).runs[0].bold = True
        for name in written:
            doc.add_paragraph(name, style="List Bullet")
    doc.add_page_break()

    for phase, names in PHASES.items():
        written = [n for n in names if sections.get(n, "").strip()]
        if not written:
            continue
        doc.add_heading(phase, level=1)
        for name in written:
            content = _display_content(sections[name].strip())
            doc.add_heading(name, level=2)
            body_lines = content.splitlines()
            if body_lines and HEADING_RE.match(body_lines[0].strip()):
                body_lines = body_lines[1:]
            render_markdown_body(doc, "\n".join(body_lines))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- voice system prompt ---------------------------------------------------

def to_voice_system_prompt(intake: BrandIntake, sections: dict[str, str]) -> str:
    """~500-token machine-usable brand voice card, meant to be dropped
    directly into another LLM's system prompt."""
    voice_section = sections.get("Brand Personality & Voice", "")
    guardrails_section = sections.get("Brand Guardrails", "")
    pillars_section = sections.get("Messaging Pillars", "")

    archetype = ARCHETYPES_BY_ID.get(intake.brand_archetype)
    archetype_line = (
        f"Brand archetype: {archetype.name} — {archetype.description}" if archetype else ""
    )

    lines = [
        f"# {intake.brand_name} — Voice Card",
        "",
        f"You are writing as {intake.brand_name}, a {intake.brand_category} brand. "
        f"One-liner: {intake.one_liner}",
        archetype_line,
        "",
        "## Tone dimensions (1=left, 7=right)",
        f"- Formal <-> Casual: {intake.personality.formal_casual}",
        f"- Serious <-> Playful: {intake.personality.serious_playful}",
        f"- Authoritative <-> Friendly: {intake.personality.authoritative_friendly}",
        f"- Classic <-> Innovative: {intake.personality.classic_innovative}",
        f"- Corporate <-> Rebellious: {intake.personality.corporate_rebellious}",
        "",
        "## Voice traits, do's and don'ts",
        voice_section.strip() or "TBD — needs founder input",
        "",
        "## Messaging pillars",
        pillars_section.strip() or "TBD — needs founder input",
        "",
        "## Guardrails (banned words, sensitive topics, competitor policy)",
        guardrails_section.strip() or "TBD — needs founder input",
        "",
        "## Never sound like",
        intake.never_sound_like,
        "",
        "Always stay within these guardrails. When information needed to answer "
        "is missing, say so plainly rather than inventing brand facts.",
    ]
    return "\n".join(lines)
