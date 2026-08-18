"""Render markdown into a python-docx document.

Extracted from app/brandforge/export.py, which had the only copy. The Marketing Plan
needs the same thing — its sections are markdown written by a model, and the user wants
them as Word documents — and a second copy would have drifted the first time one of them
learned to handle a construct the other didn't.

Deliberately not a general markdown implementation. It covers what the generators
actually emit (headings, bold, bullet and numbered lists, pipe tables) and treats
anything else as a paragraph, which is the honest failure mode: unrecognised syntax
survives as its literal text instead of vanishing.
"""
from __future__ import annotations

import re

from docx import Document

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_CELL_SPLIT_RE = re.compile(r"(?<!\\)\|")


def _fill_runs(para, line: str) -> None:
    """Write `line` into an existing paragraph, turning **bold** into real bold runs."""
    pos = 0
    for m in _BOLD_RE.finditer(line):
        if m.start() > pos:
            para.add_run(line[pos : m.start()])
        run = para.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(line):
        para.add_run(line[pos:])


def add_markdown_paragraph(doc: Document, line: str) -> None:
    _fill_runs(doc.add_paragraph(), line)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    # Split on unescaped pipes only. A plain `.split("|")` also cut on the `\|` the
    # markdown writers emit for a literal pipe, so a cell containing one still sheared the
    # row into an extra column — which is the whole thing the escaping exists to prevent.
    rows = [
        [cell.strip() for cell in _CELL_SPLIT_RE.split(line.strip().strip("|"))]
        for line in lines
        if line.strip().startswith("|")
    ]
    rows = [r for r in rows if not all(re.fullmatch(r"-+", c) for c in r)]
    if not rows:
        return
    # Ragged rows are common in model output — a row with fewer cells than the header
    # would raise on the widest row otherwise, taking the whole export down over one
    # malformed line.
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j >= width:
                break
            # Cells get the same inline treatment as body text. Assigning `.text`
            # directly was simpler and left the literal ** of every bold cell sitting in
            # the document — which plans hit constantly, since the generators bold the
            # first column of almost every table they write.
            #
            # The markdown writers also escape a literal pipe so it doesn't end the cell;
            # that escape is markdown syntax and has no business in a Word table.
            para = table.cell(i, j).paragraphs[0]
            _fill_runs(para, cell.replace(r"\|", "|"))
            if i == 0:
                for r in para.runs:
                    r.bold = True


def render_markdown_body(doc: Document, markdown: str, heading_offset: int = 0) -> None:
    """Append `markdown` to `doc`.

    `heading_offset` pushes every heading down a level, so a section can be nested under
    a document heading the caller added itself without its own `#` colliding with it.
    """
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)) + heading_offset, 9)
            doc.add_heading(heading_match.group(2), level=max(level, 1))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            i += 1
            continue

        add_markdown_paragraph(doc, stripped)
        i += 1
