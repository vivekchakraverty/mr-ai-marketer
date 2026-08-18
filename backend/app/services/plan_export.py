"""Write a generated marketing plan to disk as markdown, Word and a keyword sheet.

A plan is five things, not one: the composed strategy plus the SEO, social, paid-ads and
keyword-research work it was built from. The Space hands all five back, and the tabs in
the app show all five, so all five get exported — a bundle that only saved the composed
document would quietly drop the research the user is most likely to want in a spreadsheet
and the sections they are most likely to hand to different people.

Each aspect is written on its own *and* into one combined document, because both are
real needs: the ads plan goes to whoever buys the ads, and the whole thing goes to
whoever signs off on it.
"""
from __future__ import annotations

import csv
import io
import uuid
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .. import config
from .markdown_docx import HEADING_RE, render_markdown_body

# Order matters: it is the order of the tabs in the app, the order of the sections in the
# combined document, and the order the files are listed back to the caller.
ASPECTS: list[tuple[str, str]] = [
    ("full", "Full Plan"),
    ("keywords", "Keyword Research"),
    ("seo", "SEO Plan"),
    ("social", "Social Plan"),
    ("ads", "Ads Plan"),
]

SHEET_COLUMNS = ["Keyword", "Monthly searches", "CPC", "Related keywords", "Data source"]


@dataclass
class ExportedFile:
    """One written file, named for what it holds rather than where it landed."""

    aspect: str
    label: str
    fmt: str
    path: Path

    @property
    def url(self) -> str:
        return "/outputs/" + self.path.resolve().relative_to(config.OUTPUTS_DIR.resolve()).as_posix()


@dataclass
class PlanBundle:
    title: str
    markdown: dict[str, str]
    keyword_rows: list[dict] = field(default_factory=list)
    keyword_source_note: str = ""
    geo: str = ""
    budget_usd_per_month: float = 0
    industry_label: str = ""
    manpower_summary: str = ""


def _front_matter(bundle: PlanBundle) -> str:
    lines = [
        "---",
        f"title: {bundle.title}",
        f"date: {date.today().isoformat()}",
    ]
    if bundle.industry_label:
        lines.append(f"industry: {bundle.industry_label}")
    if bundle.geo:
        lines.append(f"geography: {bundle.geo}")
    if bundle.budget_usd_per_month:
        lines.append(f"monthly_budget_usd: {bundle.budget_usd_per_month:g}")
    if bundle.keyword_source_note:
        lines.append(f"keyword_data_source: {bundle.keyword_source_note}")
    lines.append("---")
    return "\n".join(lines) + "\n\n"


def combined_markdown(bundle: PlanBundle) -> str:
    """Every aspect in one markdown document, front matter first.

    Section headings are pushed down a level so the aspect titles added here stay above
    the model's own `#` headings instead of competing with them.
    """
    parts = [_front_matter(bundle), f"# {bundle.title}\n"]
    for key, label in ASPECTS:
        body = (bundle.markdown.get(key) or "").strip()
        if not body:
            continue
        parts.append(f"## {label}\n")
        parts.append(_demote_headings(_drop_leading_title(body, label)) + "\n")
    return "\n".join(parts)


def _drop_leading_title(markdown: str, label: str) -> str:
    """Drop the section's own opening heading when it just restates the label.

    The generators usually start a section with its own title, which is right when the
    section stands alone and redundant once it sits under a heading that already says the
    same thing — two near-identical headings back to back.
    """
    lines = markdown.splitlines()
    if not lines:
        return markdown
    match = HEADING_RE.match(lines[0].strip())
    if match and match.group(2).strip().lower() == label.strip().lower():
        return "\n".join(lines[1:]).lstrip("\n")
    return markdown


def _demote_headings(markdown: str) -> str:
    out = []
    for line in markdown.splitlines():
        match = HEADING_RE.match(line.strip())
        # Six is markdown's floor; a seventh # stops being a heading and renders as
        # literal hashes, so anything already that deep is left where it is.
        out.append("#" + line if match and len(match.group(1)) < 6 else line)
    return "\n".join(out)


# --- docx -------------------------------------------------------------------

def _docx_bytes(title: str, subtitle: str, meta: str, sections: list[tuple[str, str]]) -> bytes:
    doc = Document()

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.size = Pt(16)
    if meta:
        meta_para = doc.add_paragraph(meta)
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    multi = len(sections) > 1
    if multi:
        doc.add_page_break()
        doc.add_heading("Contents", level=1)
        for label, _ in sections:
            doc.add_paragraph(label, style="List Bullet")

    for label, body in sections:
        doc.add_page_break()
        doc.add_heading(label, level=1)
        lines = body.strip().splitlines()
        # The model usually opens a section with its own title. Kept as-is would put two
        # near-identical headings back to back, so the duplicate is dropped.
        if lines and HEADING_RE.match(lines[0].strip()):
            first = HEADING_RE.match(lines[0].strip()).group(2).strip().lower()
            if first == label.strip().lower():
                lines = lines[1:]
        render_markdown_body(doc, "\n".join(lines), heading_offset=1)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _meta_line(bundle: PlanBundle) -> str:
    bits = [b for b in (bundle.industry_label, bundle.geo.upper() if bundle.geo else "") if b]
    if bundle.budget_usd_per_month:
        bits.append(f"${bundle.budget_usd_per_month:,.0f}/month")
    bits.append(date.today().isoformat())
    return "  |  ".join(bits)


# --- keyword sheet ----------------------------------------------------------

def keyword_csv(rows: list[dict]) -> str:
    buf = io.StringIO(newline="")
    writer = csv.writer(buf)
    writer.writerow(SHEET_COLUMNS)
    for row in rows:
        writer.writerow(
            [
                row.get("keyword", ""),
                row.get("volume", ""),
                row.get("cpc", ""),
                "; ".join(row.get("related") or []),
                row.get("sourceLabel") or row.get("source", ""),
            ]
        )
    return buf.getvalue()


def _write_keyword_xlsx(path: Path, rows: list[dict]) -> bool:
    """A real spreadsheet alongside the CSV. False if openpyxl isn't available.

    Not fatal when it's missing: the CSV holds the same data and every spreadsheet
    application opens it. This is the nicer copy, not the only one.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
    except ImportError:
        return False

    wb = Workbook()
    ws = wb.active
    ws.title = "Keyword research"
    ws.append(SHEET_COLUMNS)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for row in rows:
        ws.append(
            [
                row.get("keyword", ""),
                row.get("volume", ""),
                row.get("cpc", ""),
                "; ".join(row.get("related") or []),
                row.get("sourceLabel") or row.get("source", ""),
            ]
        )
    widths = [38, 18, 16, 60, 34]
    for i, width in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = width
    ws.freeze_panes = "A2"
    wb.save(path)
    return True


# --- the bundle -------------------------------------------------------------

def run_dir() -> Path:
    d = config.OUTPUTS_DIR / "plan" / str(uuid.uuid4())
    d.mkdir(parents=True, exist_ok=True)
    return d


def _slug(text: str) -> str:
    kept = [c.lower() if c.isalnum() else "-" for c in text.strip()]
    slug = "".join(kept).strip("-")
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug[:60] or "marketing-plan"


def write_bundle(bundle: PlanBundle, directory: Path | None = None) -> list[ExportedFile]:
    """Write every aspect as .md and .docx, plus the keyword sheet, and list what landed.

    Returns the files in ASPECTS order with the combined document first, so the caller can
    render the list straight through without sorting it.
    """
    directory = directory or run_dir()
    directory.mkdir(parents=True, exist_ok=True)
    stem = _slug(bundle.title)
    meta = _meta_line(bundle)
    written: list[ExportedFile] = []

    def _write(aspect: str, label: str, fmt: str, name: str, payload: bytes | str) -> None:
        path = directory / name
        if isinstance(payload, bytes):
            path.write_bytes(payload)
        else:
            # newline="\n" so a plan exported on Windows is byte-identical to one
            # exported anywhere else, rather than silently gaining CRLF line endings.
            path.write_text(payload, encoding="utf-8", newline="\n")
        written.append(ExportedFile(aspect=aspect, label=label, fmt=fmt, path=path))

    present = [(k, l) for k, l in ASPECTS if (bundle.markdown.get(k) or "").strip()]

    # The whole plan, in both formats.
    _write("bundle", "Complete plan", "md", f"{stem}.md", combined_markdown(bundle))
    _write(
        "bundle",
        "Complete plan",
        "docx",
        f"{stem}.docx",
        _docx_bytes(
            bundle.title,
            "Marketing Plan",
            meta,
            [(label, bundle.markdown[key]) for key, label in present],
        ),
    )

    # Then each aspect on its own.
    for key, label in present:
        body = bundle.markdown[key]
        _write(key, label, "md", f"{stem}-{key}.md", body)
        _write(
            key,
            label,
            "docx",
            f"{stem}-{key}.docx",
            _docx_bytes(bundle.title, label, meta, [(label, body)]),
        )

    # The research table as data, not prose.
    if bundle.keyword_rows:
        # Written as bytes, not text. csv.writer already ends each row with CRLF, and
        # text mode on Windows translates the LF in that pair again — every row comes out
        # separated by a blank line. utf-8-sig because without the BOM Excel reads the
        # file as the local ANSI codepage and turns non-ASCII keywords into mojibake.
        path = directory / f"{stem}-keywords.csv"
        path.write_bytes(keyword_csv(bundle.keyword_rows).encode("utf-8-sig"))
        written.append(ExportedFile("keywords", "Keyword research", "csv", path))

        xlsx_path = directory / f"{stem}-keywords.xlsx"
        if _write_keyword_xlsx(xlsx_path, bundle.keyword_rows):
            written.append(ExportedFile("keywords", "Keyword research", "xlsx", xlsx_path))

    return written
